import json
import re
from typing import Any

import pandas as pd
import pdfplumber
from docx import Document

from ventilacia_ai.clients.gigachat_client import (
    GigaChatQuotaExceeded,
    get_gigachat_client,
    is_quota_exceeded,
)
from ventilacia_ai.services import config_service


def allowed_file(filename: str) -> bool:
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in config_service.ALLOWED_EXTENSIONS
    )


def extract_text_from_excel(file_path: str) -> str:
    """
    Собирает всё содержимое Excel-файла (все листы, все строки) в структурированный
    текст для передачи в LLM. Каждая строка — «Строка N: A | B | C | ...».
    """
    try:
        sheets = pd.read_excel(file_path, header=None, sheet_name=None)
    except Exception as e:
        print(f"Не удалось прочитать Excel: {e}")
        return ""

    lines: list[str] = []
    row_counter = 0
    for sheet_name, df in sheets.items():
        if len(sheets) > 1:
            lines.append(f"=== Лист: {sheet_name} ===")
        for _, row in df.iterrows():
            row_counter += 1
            cells = [
                str(v).strip() for v in row.values if pd.notna(v) and str(v).strip()
            ]
            if not cells:
                continue
            lines.append(f"Строка {row_counter}: " + " | ".join(cells))
    return "\n".join(lines)


def parse_excel_application(file_path: str) -> list[dict[str, Any]]:
    """
    Строгий парсер Excel фиксированного формата (быстро, без LLM):
    - 1 столбец: наименование
    - 2 столбец: количество
    - 3 столбец: единица измерения (опционально)
    """
    items: list[dict[str, Any]] = []
    try:
        df = pd.read_excel(file_path, header=None)

        # Набор типичных заголовков для пропуска
        header_names = {
            "наименование", "название", "name", "позиция",
            "наименование позиции", "наименование товара",
        }

        for idx, row in df.iterrows():
            if len(row) < 1:
                continue

            name = str(row.iloc[0]).strip() if pd.notna(row.iloc[0]) else ""
            if not name or name.lower() in header_names:
                continue

            quantity = "1"
            if len(row) > 1 and pd.notna(row.iloc[1]):
                qty_str = str(row.iloc[1]).strip()
                m = re.search(r"(\d+(?:[.,]\d+)?)", qty_str)
                if m:
                    quantity = m.group(1).replace(",", ".")

            unit = ""
            if len(row) > 2 and pd.notna(row.iloc[2]):
                unit_raw = str(row.iloc[2]).strip()
                if unit_raw.lower() not in {"ед.изм", "ед. изм", "ед.изм.", "единица измерения", "unit", ""}:
                    unit = unit_raw

            items.append(
                {
                    "row_number": idx + 1,
                    "name": name,
                    "quantity": quantity,
                    "unit": unit,
                }
            )

        return items
    except Exception as e:
        print(f"Ошибка при парсинге Excel: {e}")
        return []


def parse_docx_application(file_path: str) -> list[dict[str, Any]]:
    """Парсит Word документ заявки и извлекает позиции."""
    try:
        doc = Document(file_path)
        structured_text = ""

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text and len(text) > 5:
                structured_text += f"Строка {idx + 1}: {text}\n"

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    row_text = " | ".join(row_data)
                    if len(row_text) > 10:
                        structured_text += (
                            f"Таблица {table_idx + 1}, Строка {row_idx + 1}: {row_text}\n"
                        )

        return parse_text_with_ai(structured_text)
    except Exception as e:
        print(f"Ошибка при парсинге Word: {e}")
        return []


def parse_pdf_application(pdf_path: str) -> list[dict[str, Any]]:
    """Парсит PDF файл заявки и извлекает позиции."""
    try:
        table_rows = []
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:
                        for row_idx, row in enumerate(table[2:], start=3):
                            if row and len(row) > 0:
                                non_empty_cells = [
                                    str(cell).strip() if cell else ""
                                    for cell in row
                                    if cell
                                ]
                                if any(
                                    cell for cell in non_empty_cells if len(str(cell)) > 2
                                ):
                                    table_rows.append(
                                        {
                                            "page": page_num + 1,
                                            "row": row_idx,
                                            "cells": non_empty_cells,
                                        }
                                    )

        structured_text = ""
        for row_info in table_rows:
            row_text = " | ".join(row_info["cells"])
            if row_text and len(row_text) > 10:
                structured_text += f"Строка {row_info['row']}: {row_text}\n"

        full_text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"

        return parse_text_with_ai(structured_text if structured_text else full_text)
    except Exception as e:
        print(f"Ошибка при парсинге PDF: {e}")
        return []


_CHUNK_SIZE = 2500  # символов на один запрос к LLM
_MAX_RETRIES = 3    # количество попыток при сетевых/серверных сбоях


_SUPERSCRIPT_DIGITS = {
    "⁰": "0",
    "¹": "1",
    "²": "2",
    "³": "3",
    "⁴": "4",
    "⁵": "5",
    "⁶": "6",
    "⁷": "7",
    "⁸": "8",
    "⁹": "9",
}


def _normalize_degrees_and_superscripts(text: str) -> str:
    """
    Приводит разные способы записи градусов к «°» так, чтобы LLM не путала
    «90⁰» с «900». Обрабатывает надстрочные цифры, «º» (U+00BA) и литеральный
    «^0» после числа.
    """
    if not text:
        return text

    def _sup_repl(match: re.Match[str]) -> str:
        digits = "".join(_SUPERSCRIPT_DIGITS.get(ch, ch) for ch in match.group(0))
        return "°" if digits == "0" else digits

    text = re.sub(r"[⁰¹²³⁴⁵⁶⁷⁸⁹]+", _sup_repl, text)
    text = re.sub(r"(\d)\s*[ºº]", r"\1°", text)
    text = re.sub(r"(\d)\s*\^\s*0\b", r"\1°", text)
    return text


def _split_text_into_chunks(text: str, chunk_size: int = _CHUNK_SIZE) -> list[str]:
    """Разбивает текст на куски по границам строк, не превышая chunk_size."""
    if len(text) <= chunk_size:
        return [text]
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for line in text.split("\n"):
        line_len = len(line) + 1
        if current and current_len + line_len > chunk_size:
            chunks.append("\n".join(current))
            current = [line]
            current_len = line_len
        else:
            current.append(line)
            current_len += line_len
    if current:
        chunks.append("\n".join(current))
    return chunks


def parse_text_with_ai(text: str) -> list[dict[str, Any]]:
    """
    Использует GigaChat для извлечения позиций из текста.
    При большом объёме текст разбивается на части и результаты объединяются.
    """
    text = _normalize_degrees_and_superscripts(text)
    print(f"[PARSE_AI] Старт, длина текста: {len(text)} символов")
    print(f"[PARSE_AI] Превью входного текста (первые 800 симв.): {text[:800]!r}")

    chunks = _split_text_into_chunks(text)
    if len(chunks) > 1:
        print(f"[PARSE_AI] Текст разбит на {len(chunks)} частей")
        all_items: list[dict[str, Any]] = []
        for i, chunk in enumerate(chunks, 1):
            print(f"[PARSE_AI] Чанк {i}/{len(chunks)} ({len(chunk)} симв.)")
            items = _parse_chunk_with_ai(chunk)
            all_items.extend(items)
        print(f"[PARSE_AI] Итого собрано из всех чанков: {len(all_items)} позиций")
        return all_items

    return _parse_chunk_with_ai(text)


def _parse_chunk_with_ai(text: str) -> list[dict[str, Any]]:
    """Один запрос к GigaChat для части текста."""
    try:
        client = get_gigachat_client()
    except ValueError as e:
        print(f"[PARSE_AI] Не удалось создать клиент GigaChat: {e}")
        return []

    try:
        prompt = f"""Ты помощник для извлечения позиций товаров/материалов из заявки.

Из следующего текста извлеки ТОЛЬКО позиции товаров/материалов (сколько и чего клиент хочет заказать).

ВАЖНО: Игнорируй полностью:
- Заголовки документа ("Заявка на...", "Ведомость...")
- Имена людей, подписи, должности
- Даты, адреса
- Заголовки колонок таблицы ("№ п/п", "Наименование", "Ед.изм", "Количество")
- Пустые строки, разделители, сноски
- Общие примечания без конкретного количества (например: "воздуховоды круглого сечения выполнить длиной L=2м")
- Искажённый/нечитаемый текст

Правила:
- Наименование бери полностью, со всеми техническими характеристиками (ГОСТ, толщина, диаметр, размеры).
- Количество — только число (десятичный разделитель может быть "," или ".").
- Ед.изм — как в тексте: "шт", "м.", "пог.м", "м2" и т.д.
- Если единица не указана — оставь пустую строку.
- Символы Ø, ø, Ф, ф, ∅ — это ОБОЗНАЧЕНИЯ ДИАМЕТРА, сохраняй их в наименовании.
- Символ ° — это ГРАДУСЫ. Сочетания вроде "90°", "45°" НИКОГДА не приклеивай к соседним цифрам (не превращай "90°" в "900" и т.п.).

Структурированные данные:
{text}

Верни JSON массив объектов в формате:
[
    {{
        "row_number": 1,
        "name": "полное наименование с характеристиками",
        "quantity": "6",
        "unit": "м."
    }}
]"""
        full_prompt = (
            "Ты помощник для извлечения позиций товаров/материалов из заявок и ведомостей. "
            "ВАЖНО: твой ответ должен начинаться с символа '[' и заканчиваться ']'. "
            "Не пиши никаких пояснений, приветствий, заголовков или текста вне JSON-массива. "
            "Если позиций нет — верни пустой массив [].\n\n"
            + prompt
        )
        print(f"[PARSE_AI] Отправка в GigaChat, длина промпта: {len(full_prompt)}")

        response = None
        last_error: Exception | None = None
        for attempt in range(1, _MAX_RETRIES + 1):
            try:
                response = client.chat(full_prompt)
                break
            except Exception as call_err:
                # 402 ретраить бессмысленно — лимит/баланс не вернётся за пару секунд.
                if is_quota_exceeded(call_err):
                    print(
                        "[PARSE_AI] Лимит GigaChat исчерпан (HTTP 402). "
                        "Прерываю запросы без повторов."
                    )
                    raise GigaChatQuotaExceeded(
                        "Лимит или баланс GigaChat исчерпан (HTTP 402 Payment Required). "
                        "Пополните баланс в личном кабинете Сбера или дождитесь сброса лимита."
                    ) from call_err
                last_error = call_err
                print(
                    f"[PARSE_AI] Попытка {attempt}/{_MAX_RETRIES} не удалась: {call_err}"
                )
                if attempt < _MAX_RETRIES:
                    import time
                    time.sleep(2 * attempt)  # 2с, 4с, 6с

        if response is None:
            print(f"[PARSE_AI] Все попытки исчерпаны. Последняя ошибка: {last_error}")
            return []

        print(f"[PARSE_AI] Получен ответ типа: {type(response).__name__}")

        if hasattr(response, "choices") and len(response.choices) > 0:
            result_text = response.choices[0].message.content.strip()
        elif hasattr(response, "content"):
            result_text = response.content.strip()
        else:
            result_text = str(response).strip()

        print(f"[PARSE_AI] Длина ответа: {len(result_text)} символов")
        print(f"[PARSE_AI] Превью ответа: {result_text[:300]!r}")

        if not result_text:
            print("[PARSE_AI] Пустой ответ от GigaChat (возможно, исчерпан лимит)")
            return []

        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        result_text = result_text.strip()

        json_payload = _extract_json_array(result_text)
        if json_payload is None:
            preview = result_text[:500].replace("\n", " ")
            print(f"[PARSE_AI] В ответе не найден JSON-массив. Превью ответа: {preview!r}")
            return []

        try:
            items = json.loads(json_payload)
        except json.JSONDecodeError as e:
            preview = json_payload[:500].replace("\n", " ")
            print(f"[PARSE_AI] Не удалось распарсить JSON: {e}. Превью: {preview!r}")
            return []

        print(f"[PARSE_AI] JSON распарсен, позиций: {len(items) if isinstance(items, list) else 0}")
        cleaned = validate_and_clean_items(items)
        print(f"[PARSE_AI] После валидации осталось: {len(cleaned)} позиций")
        return cleaned
    except Exception as e:
        print(f"Ошибка при парсинге через нейросеть: {e}")
        return []


def _extract_json_array(text: str) -> str | None:
    """
    Извлекает первый JSON-массив верхнего уровня из текста. Полезно, когда LLM
    оборачивает JSON в пояснения на естественном языке или добавляет вступление.
    """
    if not text:
        return None
    start = text.find("[")
    if start == -1:
        return None
    depth = 0
    in_string = False
    escape = False
    for i in range(start, len(text)):
        ch = text[i]
        if escape:
            escape = False
            continue
        if ch == "\\":
            escape = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == "[":
            depth += 1
        elif ch == "]":
            depth -= 1
            if depth == 0:
                return text[start : i + 1]
    return None


def validate_and_clean_items(items: Any) -> list[dict[str, Any]]:
    """Валидация и очистка извлеченных позиций."""
    validated_items: list[dict[str, Any]] = []

    exclude_keywords = [
        "комплектовочная",
        "ведомость",
        "материалов",
        "оборудования",
        "утверждаю",
        "начальник",
        "филиала",
        "подпись",
        "инициалы",
        "фамилия",
        "января",
        "февраля",
        "марта",
        "апреля",
        "мая",
        "июня",
        "июля",
        "августа",
        "сентября",
        "октября",
        "ноября",
        "декабря",
        "года",
        "г.",
        "наименование",
        "технические",
        "характеристики",
        "тип",
        "марка",
        "гост",
        "ту",
        "завод",
        "изготовитель",
        "ед",
        "изм",
        "кол-во",
        "стоимость",
        "смете",
        "цена",
        "ндс",
        "статья",
        "затрат",
        "месяц",
        "поставки",
        "гпр",
        "примечание",
        "не предъявлено",
        "заказчику",
        "шифр",
        "объекта",
        "строительства",
        "локальный",
        "сметный",
        "расчет",
        "отсутствует",
    ]

    def is_readable_text(text: str) -> bool:
        if not text or len(text) < 4:
            return False
        has_cyrillic = bool(re.search(r"[А-Яа-яЁё]", text))
        # 6+ согласных подряд — признак OCR-мусора. 5 бывает в обычных словах.
        if re.search(r"[бвгджзклмнпрстфхцчшщБВГДЖЗКЛМНПРСТФХЦЧШЩ]{6,}", text):
            return False
        words = re.findall(r"[А-Яа-яЁё]{3,}", text)
        has_tech_info = bool(
            re.search(r"\d", text)
            or re.search(r"[øØфФΦ∅]", text)
            or re.search(r"\b(мм|м|шт|кг|гост|ту|пог)\b", text.lower())
        )
        # Минимум одно осмысленное слово + либо второе слово, либо технические характеристики.
        if has_cyrillic and len(words) == 0:
            return False
        if len(words) < 2 and not has_tech_info:
            return False
        return True

    if not isinstance(items, list):
        return []

    for item in items:
        if not isinstance(item, dict) or "name" not in item:
            continue

        name = str(item.get("name", "")).strip()
        row_number = item.get("row_number", None)

        if not name or len(name) < 5:
            continue
        if not is_readable_text(name):
            continue

        name_lower = name.lower()
        if any(keyword in name_lower for keyword in exclude_keywords):
            has_technical_info = any(char.isdigit() for char in name) and (
                "мм" in name_lower
                or "м.п" in name_lower
                or "м " in name_lower
                or "кг" in name_lower
                or "шт" in name_lower
                or "гост" in name_lower
                or "ту" in name_lower
                or len(name) > 20
            )
            if not has_technical_info:
                continue

        if re.match(r"^[А-ЯЁ]\.\s*[А-ЯЁ]\.\s*[А-ЯЁ][а-яё]+$", name):
            continue

        if re.search(
            r"\d{1,2}\s*(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s*\d{4}",
            name_lower,
        ):
            continue

        name = re.sub(r"^[\|\s]+", "", name)

        quantity = str(item.get("quantity", "1")).strip() or "1"
        quantity = re.sub(r"[^\d.,]", "", quantity)
        if not quantity or quantity == "0":
            quantity = "1"

        unit = str(item.get("unit", "")).strip()

        validated_items.append(
            {
                "row_number": row_number if row_number else len(validated_items) + 1,
                "name": name,
                "quantity": quantity,
                "unit": unit,
            }
        )

    validated_items.sort(key=lambda x: x.get("row_number", 999))
    return validated_items


def parse_application_file(file_path: str, extension: str) -> list[dict[str, Any]]:
    """Строгий режим: быстрый парсер Excel фиксированного формата, без LLM.

    Поддерживает только xlsx/xls:
    - 1 колонка: наименование
    - 2 колонка: количество
    - 3 колонка: ед.изм (опционально)
    """
    ext = (extension or "").lower().lstrip(".")
    if ext in {"xlsx", "xls"}:
        return parse_excel_application(file_path)
    return []


def parse_application_file_smart(file_path: str, extension: str) -> list[dict[str, Any]]:
    """Умный режим: универсальный парсер заявки через LLM.

    Собирает всё содержимое файла в текст и передаёт в GigaChat, который сам
    определяет, где позиции, где шапка/примечания, где имена людей и т.д.

    Поддерживаемые форматы: xlsx, xls, docx, pdf (с текстовым слоем).
    """
    ext = (extension or "").lower().lstrip(".")
    if ext in {"xlsx", "xls"}:
        text = extract_text_from_excel(file_path)
        if not text:
            return []
        return parse_text_with_ai(text)
    if ext == "docx":
        return parse_docx_application(file_path)
    if ext == "pdf":
        return parse_pdf_application(file_path)
    return []

